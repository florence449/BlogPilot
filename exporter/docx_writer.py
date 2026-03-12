import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import OUTPUT_DIR
from logger import get_logger
from exporter.parser import parse_markdown

logger = get_logger(__name__)


# ========================
# 스타일 설정
# ========================

def _set_heading_style(paragraph, level: int):
    """소제목 폰트 크기 및 색상 설정"""
    size_map = {1: 18, 2: 15, 3: 13}
    color_map = {
        1: RGBColor(0x22, 0x22, 0x22),
        2: RGBColor(0x33, 0x33, 0x33),
        3: RGBColor(0x44, 0x44, 0x44),
    }
    for run in paragraph.runs:
        run.font.size = Pt(size_map.get(level, 12))
        run.font.bold = True
        run.font.color.rgb = color_map.get(level, RGBColor(0, 0, 0))


def _add_runs(paragraph, runs: list):
    """run 리스트를 paragraph에 추가"""
    for run_data in runs:
        run = paragraph.add_run(run_data["text"])
        run.bold = run_data["bold"]
        run.font.size = Pt(11)


# ========================
# 블록 렌더러
# ========================

def _render_blocks(doc: Document, blocks: list):
    """파싱된 블록을 순서대로 docx에 렌더링"""
    for block in blocks:
        btype = block["type"]

        if btype == "empty":
            doc.add_paragraph()

        elif btype == "h1":
            p = doc.add_heading(level=1)
            p.clear()
            _add_runs(p, block["runs"])
            _set_heading_style(p, 1)

        elif btype == "h2":
            p = doc.add_heading(level=2)
            p.clear()
            _add_runs(p, block["runs"])
            _set_heading_style(p, 2)

        elif btype == "h3":
            p = doc.add_heading(level=3)
            p.clear()
            _add_runs(p, block["runs"])
            _set_heading_style(p, 3)

        elif btype == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, block["runs"])

        elif btype == "hashtag":
            p = doc.add_paragraph()
            run = p.add_run(block["raw"])
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x11, 0x55, 0xCC)

        else:  # p
            p = doc.add_paragraph()
            _add_runs(p, block["runs"])


# ========================
# docx 저장
# ========================

def save_docx(content: str, topic: str, blog_code: str) -> str:
    """
    마크다운 텍스트를 파싱해 docx로 저장
    반환: 저장된 파일 경로
    """
    logger.info(f"docx 저장 시작 - 주제: {topic}")

    blocks = parse_markdown(content)

    doc = Document()

    # 기본 폰트 설정
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    _render_blocks(doc, blocks)

    # 파일명 생성
    today = datetime.now().strftime("%Y%m%d_%H%M")
    safe_topic = re.sub(r'[^\w가-힣]', '_', topic)[:20]
    filename = f"{today}_{blog_code}_{safe_topic}.docx"
    filepath = OUTPUT_DIR / filename

    doc.save(str(filepath))
    logger.info(f"docx 저장 완료: {filename}")

    return str(filepath)